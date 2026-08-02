import os
import uuid
import json
import shutil
import logging
from typing import List, Optional
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from openai import OpenAI
from dotenv import load_dotenv
import docx

# Import our custom NLP and parser utils
from nlp_utils import compute_nlp_similarities
from parser_utils import parse_resume

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load env variables
load_dotenv()

app = FastAPI(title="AI Resume Screening Agent")

# Enable CORS for local development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Setup paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMP_DIR = os.path.join(BASE_DIR, "temp")
SAMPLES_DIR = os.path.join(BASE_DIR, "sample_resumes")
FRONTEND_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "frontend"))

os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(SAMPLES_DIR, exist_ok=True)
os.makedirs(FRONTEND_DIR, exist_ok=True)

# Initialize OpenAI Client
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key or openai_api_key.strip() == "" or openai_api_key == "your-api-key-here":
    logger.warning("No valid OPENAI_API_KEY found. Running in CLASSICAL NLP similarity mode.")
    client = None
    is_mock_mode = True
else:
    client = OpenAI(api_key=openai_api_key)
    is_mock_mode = False

# API Status check
@app.get("/api/status")
async def get_status():
    return {
        "status": "ok",
        "openai_configured": not is_mock_mode,
        "message": "AI Resume Screening Agent is active."
    }

# Helper to run LLM Resume Analysis
def run_llm_analysis(resume_text: str, job_description: str) -> dict:
    """
    Calls OpenAI to extract resume details and perform structured gap analysis.
    """
    if is_mock_mode or not client:
        return run_local_rule_analysis(resume_text, job_description)

    try:
        prompt = (
            "You are an expert technical recruiter and HR analyst. "
            "Analyze the following resume text against the Job Description. "
            "Extract details and evaluate strengths, weaknesses, gaps, and assign a score (0 to 100) based on alignment.\n\n"
            f"=== Job Description ===\n{job_description}\n\n"
            f"=== Resume Text ===\n{resume_text}\n\n"
            "Return a valid JSON object matching this EXACT schema:\n"
            "{\n"
            '  "candidate_name": "Full name of candidate",\n'
            '  "skills": ["Extracted skills relevant to the role"],\n'
            '  "experience": "Brief summary of candidate\'s experience",\n'
            '  "education": "Highest degree / school",\n'
            '  "strengths": ["Key strength 1", "Key strength 2"],\n'
            '  "weaknesses": ["Key weakness 1", "Key weakness 2"],\n'
            '  "gaps": ["Gaps or missing requirements"],\n'
            '  "reasoning": "Detailed, professional paragraph explaining the score evaluation",\n'
            '  "score": 85\n'
            "}\n"
            "Do not include any markdown wrappers or text outside the JSON object."
        )

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You are a professional HR assistant that outputs strict JSON formats."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=1000
        )
        
        result_text = response.choices[0].message.content
        return json.loads(result_text)

    except Exception as e:
        logger.exception("OpenAI analysis failed. Falling back to local analysis.")
        return run_local_rule_analysis(resume_text, job_description)

# Local keyword/regex parsing fallback (when OpenAI is unavailable)
def run_local_rule_analysis(resume_text: str, job_description: str) -> dict:
    """
    Fall back analysis using keyword checking and heuristics.
    """
    lines = resume_text.splitlines()
    candidate_name = "Unknown Candidate"
    
    # Try to find name in first few non-empty lines
    for line in lines[:5]:
        clean = line.strip()
        if len(clean) > 3 and not any(keyword in clean.lower() for keyword in ["resume", "curriculum", "cv", "contact", "email", "phone"]):
            candidate_name = clean
            break
            
    # Try to extract common skills using simple dictionary matching
    common_skills = [
        "python", "javascript", "react", "fastapi", "django", "flask", "postgresql", "sql", "aws", "docker", 
        "kubernetes", "git", "java", "c++", "html", "css", "nodejs", "typescript", "figma", "agile", "scrum", "jira"
    ]
    extracted_skills = []
    resume_lower = resume_text.lower()
    for skill in common_skills:
        if re_search := r"\b" + skill + r"\b":
            import re
            if re.search(re_search, resume_lower):
                extracted_skills.append(skill.capitalize())
                
    # Detect Education
    education = "Not specified"
    edu_keywords = ["bachelor", "master", "phd", "b.s", "m.s", "degree", "university", "college", "btech", "mtech"]
    for line in lines:
        if any(keyword in line.lower() for keyword in edu_keywords):
            education = line.strip()
            if len(education) > 80:
                education = education[:77] + "..."
            break

    # Analyze overlaps for strengths/weaknesses
    jd_words = set(job_description.lower().split())
    res_words = set(resume_lower.split())
    matches = jd_words.intersection(res_words)
    
    strengths = [f"Found relevant keywords: {', '.join(list(matches)[:4])}"] if matches else ["Basic keywords check completed."]
    gaps = ["OpenAI API key was not configured; deep requirements check skipped."]
    weaknesses = ["Detailed semantic analysis skipped."]
    
    return {
        "candidate_name": candidate_name,
        "skills": extracted_skills if extracted_skills else ["Text parsed successfully"],
        "experience": "Successfully extracted candidate profile.",
        "education": education,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "gaps": gaps,
        "reasoning": "This assessment was computed using localized TF-IDF and keyword statistics because the OpenAI API key is missing or invalid.",
        "score": 50 # Default baseline
    }

# Process screening core logic
def process_resumes(resume_paths: List[str], job_description: str) -> List[dict]:
    # 1. Parse text from all files
    resumes_text = []
    candidates_meta = []
    
    for path in resume_paths:
        try:
            text = parse_resume(path)
            resumes_text.append(text)
            candidates_meta.append({
                "filename": os.path.basename(path),
                "filepath": path,
                "text": text
            })
        except Exception as e:
            logger.error(f"Failed to parse {path}: {str(e)}")
            # Keep placeholder to keep indices aligned
            resumes_text.append("")
            candidates_meta.append({
                "filename": os.path.basename(path),
                "filepath": path,
                "text": "",
                "error": str(e)
            })

    # 2. Compute Cosine Similarities via our custom NLP engine
    nlp_scores = compute_nlp_similarities(job_description, resumes_text)
    
    # 3. Compute LLM Evaluations and merge scores
    results = []
    for i, meta in enumerate(candidates_meta):
        if meta.get("text") == "":
            results.append({
                "filename": meta["filename"],
                "candidate_name": meta["filename"],
                "nlp_score": 0.0,
                "llm_score": 0.0,
                "hybrid_score": 0.0,
                "skills": [],
                "experience": "Failed to parse document text.",
                "education": "N/A",
                "strengths": [],
                "weaknesses": ["Could not extract text from document"],
                "gaps": [meta.get("error", "Unknown parsing error")],
                "reasoning": "This candidate was rejected automatically due to document structure extraction failure."
            })
            continue

        # Get LLM evaluation details
        eval_data = run_llm_analysis(meta["text"], job_description)
        
        # Calculate final hybrid score
        # Cosine similarity is usually 0.0 to 0.5 for real text. Let's scale NLP score to 0-100 range.
        nlp_score_scaled = min(nlp_scores[i] * 2.0 * 100.0, 100.0) # Cosine similarity of 0.5 becomes 100.
        llm_score = float(eval_data.get("score", 50))
        
        # Hybrid formula
        hybrid_score = (0.3 * nlp_score_scaled) + (0.7 * llm_score)
        
        results.append({
            "filename": meta["filename"],
            "candidate_name": eval_data.get("candidate_name", "Unknown Candidate"),
            "nlp_score": round(nlp_scores[i] * 100.0, 2), # Original percentage
            "llm_score": round(llm_score, 2),
            "hybrid_score": round(hybrid_score, 2),
            "skills": eval_data.get("skills", []),
            "experience": eval_data.get("experience", ""),
            "education": eval_data.get("education", ""),
            "strengths": eval_data.get("strengths", []),
            "weaknesses": eval_data.get("weaknesses", []),
            "gaps": eval_data.get("gaps", []),
            "reasoning": eval_data.get("reasoning", "")
        })
        
    # Sort candidates by hybrid score descending
    results.sort(key=lambda x: x["hybrid_score"], reverse=True)
    return results

@app.post("/api/screen")
async def screen_resumes(
    job_description: str = Form(...),
    files: List[UploadFile] = File(...)
):
    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="No resume files uploaded.")
    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description cannot be empty.")

    # Save uploaded files temporarily
    saved_paths = []
    for file in files:
        file_id = str(uuid.uuid4())
        _, ext = os.path.splitext(file.filename)
        temp_filename = f"{file_id}{ext}"
        temp_path = os.path.join(TEMP_DIR, temp_filename)
        
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        saved_paths.append(temp_path)
        
    try:
        results = process_resumes(saved_paths, job_description)
        return {
            "success": True,
            "mode": "hybrid" if not is_mock_mode else "classical-nlp",
            "candidates": results
        }
    finally:
        # Clean up temp files
        for path in saved_paths:
            if os.path.exists(path):
                os.remove(path)

# Screen the generated test sample resumes directly
@app.post("/api/screen-samples")
async def screen_samples(job_description: str = Form(...)):
    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description cannot be empty.")
        
    # Check if sample files are generated, if not, generate them first
    sample_files = [f for f in os.listdir(SAMPLES_DIR) if os.path.isfile(os.path.join(SAMPLES_DIR, f))]
    if len(sample_files) < 10:
        generate_sample_resumes_files()
        sample_files = [f for f in os.listdir(SAMPLES_DIR) if os.path.isfile(os.path.join(SAMPLES_DIR, f))]

    sample_paths = [os.path.join(SAMPLES_DIR, f) for f in sample_files]
    
    results = process_resumes(sample_paths, job_description)
    return {
        "success": True,
        "mode": "hybrid" if not is_mock_mode else "classical-nlp",
        "candidates": results
    }

# Seeder endpoint to generate 10+ sample resumes on the disk
@app.post("/api/generate-samples")
async def generate_samples():
    try:
        generate_sample_resumes_files()
        files_created = os.listdir(SAMPLES_DIR)
        return {
            "success": True,
            "message": f"Successfully generated {len(files_created)} sample resumes.",
            "files": files_created
        }
    except Exception as e:
        logger.exception("Failed to generate sample files")
        raise HTTPException(status_code=500, detail=f"Failed to generate files: {str(e)}")

def generate_sample_resumes_files():
    """
    Creates 10 mock resume documents (TXT and DOCX) in SAMPLES_DIR representing different profiles.
    """
    # Clean SAMPLES_DIR first
    for f in os.listdir(SAMPLES_DIR):
        path = os.path.join(SAMPLES_DIR, f)
        if os.path.isfile(path):
            os.remove(path)

    profiles = [
        {
            "name": "Sarah_Jenkins_Senior_Python_Developer.docx",
            "is_docx": True,
            "content": {
                "name": "Sarah Jenkins",
                "email": "sarah.jenkins@devmail.com",
                "phone": "+1 (555) 019-2834",
                "summary": "Highly motivated Senior Software Engineer with 6+ years of experience specializing in Python, FastAPI, Django, and database management using PostgreSQL. Proven track record of developing scalable APIs, microservices, and leading backend engineering teams.",
                "experience": [
                    "Lead Backend Engineer at CloudScale Solutions (2022 - Present):\n- Designed and implemented 15+ REST APIs using FastAPI and Pydantic, improving latency by 35%.\n- Managed Postgres database clustering, query optimizations, and index tune-ups.\n- Standardized deployment workflows with Docker and GitLab CI/CD.",
                    "Senior Software Developer at PyTech Corp (2020 - 2022):\n- Built high-traffic web applications with Django and PostgreSQL.\n- Collaborated with frontend teams to integrate React web apps via secure OAuth2 tokens."
                ],
                "skills": ["Python", "FastAPI", "Django", "PostgreSQL", "SQL", "Docker", "REST APIs", "AWS", "Git"],
                "education": "B.S. in Computer Science - University of Texas at Austin (2019)"
            }
        },
        {
            "name": "Amanda_White_Frontend_React_Developer.txt",
            "is_docx": False,
            "text": """AMANDA WHITE
Email: amanda.white@webdev.org | Phone: +1 (555) 012-9988
Website: amandacodes.io

SUMMARY:
Passionate Frontend Engineer with 4 years of experience building modern, highly-responsive web applications using React, JavaScript, HTML5, CSS3, Tailwind CSS, and Vite. Dedicated to pixel-perfect layouts and smooth user experience animations.

TECHNICAL SKILLS:
React, React Hooks, JavaScript (ES6+), TypeScript, HTML5, CSS3, Tailwind CSS, Sass, Vite, Git, Jest, Responsive Web Design

EXPERIENCE:
Frontend Developer | PixelCraft Studio (2022 - Present)
- Designed and built responsive landing pages and dashboard modules in React.
- Used Tailwind CSS to create custom utility tokens and implement dark-mode theme transitions.
- Reduced webpack bundle size by 40% by migrating projects to Vite.

Web UI Engineer | Interface Hub (2020 - 2022)
- Maintained legacy React and Vue platforms, translating PSD files into functional HTML/CSS templates.
- Conducted cross-browser compatibility testing and resolved accessibility issues.

EDUCATION:
B.A. in Web Design & Interactive Media - Boston University (2020)"""
        },
        {
            "name": "Sophia_Martinez_Technical_Product_Manager.docx",
            "is_docx": True,
            "content": {
                "name": "Sophia Martinez",
                "email": "sophia.martinez@pmhub.com",
                "phone": "+1 (555) 443-8822",
                "summary": "Technical Product Manager with 5+ years of experience leading cross-functional teams of engineers and designers to launch SaaS products. Expert in roadmap prioritization, customer research, Agile methodologies, and metrics analysis.",
                "experience": [
                    "Product Manager at Flowchart SaaS (2021 - Present):\n- Directed the product life cycle for enterprise analytics features from concept to launch.\n- Conducted user interviews and prioritized product backlog using Jira and Productboard.\n- Collaborated with tech leads to translate specifications into clear developer tickets.",
                    "Associate PM at DataPulse (2019 - 2021):\n- Analyzed SQL queries and data dashboards to define user activation rate milestones."
                ],
                "skills": ["Product Management", "Agile Roadmap", "Scrum", "Jira", "JQL / SQL", "Wireframing", "Figma", "User Analytics"],
                "education": "MBA - Harvard Business School (2019), B.S. in Business Information Systems - UCLA (2017)"
            }
        },
        {
            "name": "Alex_Rivera_Full_Stack_Node_Developer.txt",
            "is_docx": False,
            "text": """ALEX RIVERA
Email: alex.rivera@stackdev.io | Phone: +1 (555) 332-9011

SUMMARY:
Versatile Full Stack Software Engineer with 5 years of professional experience writing Javascript/Typescript. Expert in Node.js, Express, React, and MongoDB database environments. Strong knowledge of system architectures and REST API integrations.

CORE SKILLS:
Node.js, Express, React, Javascript, Typescript, MongoDB, SQL, Git, REST APIs, GraphQL, Heroku, Docker

PROFESSIONAL EXPERIENCE:
Full Stack Developer | AppForge Co (2021 - Present)
- Built enterprise Node.js APIs utilizing Express, handling over 100k daily requests.
- Developed dynamic user dashboards in React, integrating real-time charts and websockets.
- Leveraged MongoDB aggregation queries to build customized admin reporting modules.

Software Developer | DevEngine (2019 - 2021)
- Developed and maintained server-side Javascript modules and SQL queries.
- Created microservices to automate automated email broadcasts and slack alerts.

EDUCATION:
B.S. in Computer Science - University of Illinois (2019)"""
        },
        {
            "name": "Emily_Chen_Machine_Learning_Engineer.docx",
            "is_docx": True,
            "content": {
                "name": "Emily Chen",
                "email": "emily.chen@ai-labs.co",
                "phone": "+1 (555) 887-1234",
                "summary": "Machine Learning Research Engineer with a focus on Natural Language Processing (NLP) and computer vision. Strong academic background and 3+ years of commercial experience implementing PyTorch and Transformers model training pipelines.",
                "experience": [
                    "AI Researcher at DeepText Labs (2022 - Present):\n- Finetuned LLMs and Bert architectures for sentiment analysis and text summarization.\n- Engineered custom preprocessing pipelines for unstructured text data, parsing PDFs and HTML corpora.\n- Deployed models as lightweight FastAPI endpoints containerized on AWS.",
                    "ML Engineer at VisionAI (2020 - 2022):\n- Developed object-detection neural networks using PyTorch and OpenCV."
                ],
                "skills": ["Python", "PyTorch", "TensorFlow", "FastAPI", "NLP", "LLMs", "Scikit-Learn", "AWS", "Git"],
                "education": "M.S. in Artificial Intelligence - Stanford University (2020)"
            }
        },
        {
            "name": "Michael_Chang_DevOps_Engineer.txt",
            "is_docx": False,
            "text": """MICHAEL CHANG
Email: michael.chang@cloudops.net | Phone: +1 (555) 776-5432

SUMMARY:
Site Reliability and DevOps Engineer with 7 years of experience automating infrastructure, scaling cloud architectures, and designing robust CI/CD pipelines. Certified Kubernetes Administrator.

TECHNICAL SKILLS:
AWS, Terraform, Kubernetes, Docker, Jenkins, Git, CI/CD pipelines, Bash, Python, Linux Systems

EXPERIENCE:
DevOps Team Lead | CloudStream Tech (2021 - Present)
- Provisioned secure, multi-region AWS cloud architectures using Terraform files.
- Orchestrated container deployment across Kubernetes clusters, reducing hosting bills by 20%.
- Established automated testing and dockerization deployment scripts using Jenkins and Git hooks.

System Administrator | NetGate Solutions (2018 - 2021)
- Managed Linux system environments, handled automated backup routines, and monitored server loads.

EDUCATION:
B.S. in Computer Engineering - Georgia Tech (2018)"""
        },
        {
            "name": "Jessica_Taylor_UX_UI_Designer.docx",
            "is_docx": True,
            "content": {
                "name": "Jessica Taylor",
                "email": "jessica.taylor@designworks.com",
                "phone": "+1 (555) 909-0808",
                "summary": "Creative UX/UI Designer with 4 years of experience crafting interactive prototypes, wireframes, and design systems. Highly proficient in user-centered design, prototyping, and collaborating with developer teams.",
                "experience": [
                    "Product Designer at VisualFlow Studio (2022 - Present):\n- Designed 12 interactive prototypes in Figma, validating layouts through user usability testing.\n- Created a cohesive design system mapping theme tokens, typography, and buttons, increasing component reuse.\n- Handed off mockups to React developers with fully detailed specs.",
                    "UI Designer at BrandCraft (2020 - 2022):\n- Structured branding assets, created vector illustrations, and styled responsive landing layouts."
                ],
                "skills": ["Figma", "Adobe XD", "Wireframing", "User Research", "Prototyping", "Design Systems", "HTML/CSS Basics"],
                "education": "BFA in Graphic Design - Rhode Island School of Design (RISD) (2020)"
            }
        },
        {
            "name": "James_Wilson_Junior_Backend_Developer.docx",
            "is_docx": True,
            "content": {
                "name": "James Wilson",
                "email": "james.wilson@juniorcode.dev",
                "phone": "+1 (555) 765-4321",
                "summary": "Enthusiastic Junior Software Engineer with a solid foundation in Python, SQL databases, and web frameworks like Flask and Django. Eager to contribute to backend systems and API optimization tasks.",
                "experience": [
                    "Junior Developer Intern at ByteCodes (2023 - 2024):\n- Assisted in writing database scripts and schemas for PostgreSQL.\n- Maintained unit tests for Flask REST API endpoints, improving code coverage by 15%.\n- Debugged software configurations under senior engineer supervision."
                ],
                "skills": ["Python", "Flask", "Django", "SQL", "PostgreSQL", "Git", "REST APIs", "FastAPI basics"],
                "education": "B.S. in Software Engineering - Arizona State University (2023)"
            }
        },
        {
            "name": "Ryan_Patel_Data_Analyst.docx",
            "is_docx": True,
            "content": {
                "name": "Ryan Patel",
                "email": "ryan.patel@datawise.org",
                "phone": "+1 (555) 234-5678",
                "summary": "Detail-oriented Data Analyst with 3 years of experience converting complex transactional datasets into actionable business intelligence. Expert in SQL, Python pandas, and building interactive Tableau reports.",
                "experience": [
                    "Data Analyst at RetailGroup Inc (2022 - Present):\n- Wrote complex SQL queries in PostgreSQL to analyze weekly customer retention rates.\n- Created automated Python scripts to scrub and format messy supplier spreadsheets.\n- Built sales dashboard templates in Tableau, presented monthly to directors.",
                    "Junior Business Analyst at TechCorp (2021 - 2022):\n- Analyzed marketing funnel data using Excel pivots and SQL queries."
                ],
                "skills": ["SQL", "PostgreSQL", "Python", "Pandas / Numpy", "Tableau", "Excel", "Data Cleansing", "Data Visualization"],
                "education": "B.S. in Statistics - Ohio State University (2021)"
            }
        },
        {
            "name": "David_Miller_Agile_Project_Manager.txt",
            "is_docx": False,
            "text": """DAVID MILLER
Email: david.miller@pmgroup.com | Phone: +1 (555) 898-7654

SUMMARY:
Experienced Project Manager and Scrum Master with 8 years of experience managing software development lifecycle delivery. Strong practitioner of Agile methodologies, prioritizing sprint scope, and eliminating team blockers.

CORE SKILLS:
Agile Methodologies, Scrum, Kanban, Jira, Confluence, Risk Management, Resource Allocation, Stakeholder Communication

EXPERIENCE:
Senior Project Manager | TechFlow Inc (2020 - Present)
- Facilitated daily standups, sprint planning, and retrospective sessions for 3 software engineering squads.
- Implemented sprint velocity tracking tools in Jira, leading to 25% improvement in deployment predictability.
- Managed vendor relations and reported project status reports to executive directors.

Project Manager | BuildIT Software (2018 - 2020)
- Guided engineering groups through transition from waterfall cycles into Agile scrum methodologies.
- Coordinated user testing feedback sessions and mapped feature validation timelines.

EDUCATION:
PMP Certified (2020), B.S. in Project Management - University of Washington (2017)"""
        }
    ]

    for p in profiles:
        file_path = os.path.join(SAMPLES_DIR, p["name"])
        
        if p["is_docx"]:
            # Write a docx using python-docx
            doc = docx.Document()
            c = p["content"]
            doc.add_heading(c["name"], 0)
            doc.add_paragraph(f"Email: {c['email']} | Phone: {c['phone']}")
            
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(c["summary"])
            
            doc.add_heading("Professional Experience", level=1)
            for exp in c["experience"]:
                doc.add_paragraph(exp)
                
            doc.add_heading("Technical Skills", level=1)
            doc.add_paragraph(", ".join(c["skills"]))
            
            doc.add_heading("Education", level=1)
            doc.add_paragraph(c["education"])
            
            doc.save(file_path)
        else:
            # Write plain text
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(p["text"])
                
    logger.info("Sample resumes seeded on local storage.")

# Mount the static files from frontend.
# FastAPI checks routes in order, so mount it LAST.
if os.path.exists(FRONTEND_DIR) and os.path.exists(os.path.join(FRONTEND_DIR, "index.html")):
    app.mount("/", StaticFiles(directory=FRONTEND_DIR, html=True), name="frontend")
else:
    @app.get("/")
    async def home_fallback():
        return {
            "status": "warning",
            "message": "Frontend static files directory is missing or empty. Please create frontend/index.html."
        }

if __name__ == "__main__":
    import uvicorn
    # Seed sample resumes on launch automatically
    try:
        generate_sample_resumes_files()
    except Exception as e:
        logger.warning(f"Auto-seeding of sample resumes failed: {str(e)}")
        
    logger.info("Starting FastAPI Server...")
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
