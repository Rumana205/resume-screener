import os
import pypdf
import docx
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts plain text from a PDF file using pypdf.
    """
    text = ""
    try:
        reader = pypdf.PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Error parsing PDF file {file_path}: {str(e)}")
        raise RuntimeError(f"Could not parse PDF: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts plain text from a Word DOCX file using python-docx.
    """
    try:
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        
        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
                    
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
        raise RuntimeError(f"Could not parse DOCX: {str(e)}")

def extract_text_from_txt(file_path: str) -> str:
    """
    Extracts plain text from a TXT file with multiple encoding fallbacks.
    """
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "utf-16"]
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
    
    # Absolute fallback
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as e:
        logger.error(f"Error parsing TXT file {file_path}: {str(e)}")
        raise RuntimeError(f"Could not parse TXT: {str(e)}")

def parse_resume(file_path: str) -> str:
    """
    Inspects file extension and routes it to the appropriate parser.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".txt", ".md"]:
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported formats are .pdf, .docx, .txt")
